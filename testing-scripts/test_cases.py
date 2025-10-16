#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from datetime import datetime

def get_driver():
    if not hasattr(get_driver, 'driver'):
        options = webdriver.ChromeOptions()
        options.add_argument("--headless")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        get_driver.driver = webdriver.Chrome(
            service=Service(ChromeDriverManager().install()), 
            options=options
        )
    return get_driver.driver

BASE_DIR   = os.path.dirname(os.path.abspath(__file__))
REPORT_DOC = os.path.join(BASE_DIR, "reports", "testcases_report.docx")

def do_login(driver, base_url, email, password, timeout=10):
    login_url = f"{base_url}/admin/login"
    driver.get(login_url)
    WebDriverWait(driver, timeout).until(
        EC.presence_of_element_located((By.NAME, "email"))
    )
    driver.find_element(By.NAME, "email").send_keys(email)
    driver.find_element(By.NAME, "password").send_keys(password)
    try:
        driver.find_element(By.ID, "button-login").click()
    except:
        driver.find_element(By.CSS_SELECTOR, "button[type='submit']").click()
    WebDriverWait(driver, timeout).until(
        EC.url_contains("/admin/main")
    )

def run_test_cases():
    doc = Document()
    doc.add_heading("📄 Reporte de Test Cases Automatizados", level=0)
    doc.add_paragraph(f"📅 Fecha de ejecución: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("Este reporte incluye los resultados detallados de los casos de prueba ejecutados automáticamente.\n")

    # Crear tabla
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Nombre"
    hdr_cells[1].text = "Objetivo"
    hdr_cells[2].text = "Pasos"
    hdr_cells[3].text = "Esperado"
    hdr_cells[4].text = "Obtenido"
    hdr_cells[5].text = "Estado"
    hdr_cells[6].text = "Duración"

    chrome_options = Options()
    chrome_options.add_argument("--headless")
    driver = webdriver.Chrome(options=chrome_options)

    base_url = "http://127.0.0.1:8000"
    email = "admin@gmail.com"
    password = "admin123"

    # ---------- TC-01: Login ----------
    start = time.time()
    estado = ""
    obtenido = ""
    try:
        do_login(driver, base_url, email, password)
        estado = "✅ Aprobado"
        obtenido = "Dashboard cargado correctamente"
    except Exception as e:
        estado = "❌ Falló"
        obtenido = f"Error en login: {str(e)}"
    end = time.time()

    row = table.add_row().cells
    row[0].text = "TC-01: Login"
    row[1].text = "Verificar que un usuario válido puede iniciar sesión"
    row[2].text = "1. Ir a /admin/login\n2. Ingresar credenciales\n3. Click en Login"
    row[3].text = "Redirección al dashboard"
    row[4].text = obtenido
    row[5].text = estado
    row[6].text = f"{end - start:.2f}s"

    if estado == "❌ Falló":
        driver.quit()
        doc.save(REPORT_DOC)
        print(f"Finalizado con error. Reporte generado en: {REPORT_DOC}")
        return

    # ---------- TC-02: Crear Categoría ----------
    start = time.time()
    estado = ""
    obtenido = ""
    try:
        driver.get(f"{base_url}/admin/crud/create/categoria-resources")
        WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.NAME, "model[Nombre]"))
        )

        driver.find_element(By.NAME, "model[Nombre]").clear()
        driver.find_element(By.NAME, "model[Nombre]").send_keys("Categoría de Prueba")
        driver.find_element(By.NAME, "model[Descripcion]").clear()
        driver.find_element(By.NAME, "model[Descripcion]").send_keys("Descripción de prueba")
        driver.find_element(By.NAME, "model[Tipo_Material]").clear()
        driver.find_element(By.NAME, "model[Tipo_Material]").send_keys("Cerámica")
        driver.find_element(By.NAME, "model[Resistencia]").clear()
        driver.find_element(By.NAME, "model[Resistencia]").send_keys("Alta")

        WebDriverWait(driver, 5).until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, "button[form='post-form'][type='submit']"))
        )
        driver.find_element(By.CSS_SELECTOR, "button[form='post-form'][type='submit']").click()

        WebDriverWait(driver, 10).until(
            EC.url_contains("/admin/crud/list/categoria-resources")
        )

        if "Categoría de Prueba" in driver.page_source:
            estado = "✅ Aprobado"
            obtenido = "Categoría creada y listada correctamente"
        else:
            estado = "❌ Falló"
            obtenido = "Categoría no visible en la lista"
    except Exception as e:
        estado = "❌ Falló"
        obtenido = f"Error al crear categoría: {str(e)}"
    end = time.time()

    row = table.add_row().cells
    row[0].text = "TC-02: Crear Categoría"
    row[1].text = "Validar que un usuario puede crear una categoría"
    row[2].text = "1. Ir a crear categoría\n2. Llenar formulario\n3. Click en guardar\n4. Revisar lista"
    row[3].text = "Nueva categoría visible en la lista"
    row[4].text = obtenido
    row[5].text = estado
    row[6].text = f"{end - start:.2f}s"

    # Finalizar
    driver.quit()
    doc.save(REPORT_DOC)
    print(f"✅ Test cases finalizados. Reporte generado en: {REPORT_DOC}")

if __name__ == "__main__":
    run_test_cases()
