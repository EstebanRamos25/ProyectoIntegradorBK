#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from docx import Document
import os, time

BASE_DIR    = os.path.dirname(os.path.abspath(__file__))
REPORT_DOC  = os.path.join(BASE_DIR, "reports", "integration_report.docx")

def run_integration_tests():
    doc = Document()
    doc.add_heading("Reporte de Pruebas de Integración", level=0)

    # Configurar Selenium
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    driver = webdriver.Chrome(options=chrome_options)

    ### INT-01: Flujo completo usuario → producto
    doc.add_heading("INT-01: Flujo usuario → creación de producto", level=1)
    try:
        # 1) Login directo con usuario admin
        driver.get("http://localhost/orchidINT/login")
        time.sleep(1)
        driver.find_element(By.NAME, "email").send_keys("admin@example.com")
        driver.find_element(By.NAME, "password").send_keys("secret")
        driver.find_element(By.CSS_SELECTOR, "button[type='submit']").click()
        time.sleep(1)

        # 2) Crear categoría (si no existe)
        driver.get("http://localhost/orchidINT/admin/categories/create")
        time.sleep(1)
        driver.find_element(By.NAME, "name").send_keys("Integración Cat")
        driver.find_element(By.CSS_SELECTOR, "button[type='submit']").click()
        time.sleep(1)

        # 3) Crear producto asociado
        driver.get("http://localhost/orchidINT/admin/products/create")
        time.sleep(1)
        driver.find_element(By.NAME, "name").send_keys("Producto Integración")
        driver.find_element(By.NAME, "price").send_keys("9.99")
        # Asumimos que el select de categorías tiene name="category_id"
        driver.find_element(By.NAME, "category_id").send_keys("Integración Cat")
        driver.find_element(By.CSS_SELECTOR, "button[type='submit']").click()
        time.sleep(1)

        # 4) Verificar que el producto aparezca en el listado público
        driver.get("http://localhost/orchidINT/products")
        time.sleep(1)
        if "Producto Integración" in driver.page_source:
            doc.add_paragraph("Resultado: PASSED")
        else:
            doc.add_paragraph("Resultado: FAILED (Producto no encontrado en público)")
    except Exception as e:
        doc.add_paragraph("Resultado: ERROR (" + str(e) + ")")

    driver.quit()
    doc.save(REPORT_DOC)
    print(f"Pruebas de integración finalizadas. Reporte generado en: {REPORT_DOC}")

if __name__ == "__main__":
    run_integration_tests()
